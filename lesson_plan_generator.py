#!/usr/bin/env python3

import json
import os
import sys
import argparse
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Optional
from pathlib import Path
import re
try:
    from groq import Groq
except Exception:
    try:
        from groq.client import Groq
    except Exception:
        Groq = None

from docx import Document


# MANIFESTO


# Force UTF-8 output for Windows console
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except:
        pass

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "[PLZZ ENTER YOUR FREE GROQ API HERE]")
REFS_FOLDER = "refs/"
OUTPUT_FOLDER = "output/"
LP_TEMPLATE = "LP_formatted.docx"
DOWNLOAD_TEMPLATE = os.path.expanduser(r"C:\Users\lenovo\AppData\Local\Programs\Python\Python310\LESSON PLAN\LP_formatted.docx")   #<---- CHANGE THE FILE ADRESS FOR YOUR FORMAT

REFERENCE_FILES = [
    "Learning Outcome at Secondary Stage (1)-55-65 (1).txt",
    "Learning_Standards_Science.txt",
    "SAS_CBLF_Scientific-Literacy-Final.txt",
    "TeachersResource_LODoc-781-906 (2).txt",
]

LODOC_FILENAME = "TeachersResource_LODoc-781-906 (2).txt"

# Use ASCII-friendly symbols
CHECK = "[OK]"
CROSS = "[X]"
WARNING = "[!]"
ARROW = "->"

DEFAULT_LOCAL_CONTENT = {
    'OBJECTIVES': 'Student will be able to recall key concepts and apply them to simple problems.',
    'LEARNING_OUTCOMES': 'Identify, explain, and apply core concepts from the chapter in class activities.',
    'MISCONCEPTIONS': 'Confusing concept A with B; overlooking key conditions for X.',
    'PREV_KNOWLEDGE': 'Basic definitions and prior chapters covering fundamentals.',
    'ICE_BREAKER': 'Quick 2-min question relating chapter to everyday example.',
    'TEACHING_AIDS': 'Charts, labelled diagrams, projector slides, simple lab kit.',
    'QUESTION_EXIT_TICKET': '2 short questions assessing application of today\'s concept.',
    'AI_INTEGRATION': 'Use chat assistants for explanation and simulations; data-visualisation tools for graphs.',
    'HOMEWORK': 'Short worksheet + one applied research prompt.',
    'MI_MAPPING_SDG': 'Logical, Spatial; relates to SDG 4 (Quality Education).',
    'REFLECTION': 'What did you find easy/difficult? One-minute write-up.',
    'CAREER_RS_IDEALS': 'List of careers where the topic applies (e.g., engineering, healthcare).',
    'VOLUNTEERING': 'Community activity linking chapter concepts to local issues.',
    'TOPIC': 'Period-wise topics as provided by the teacher.',
}


# REFERENCE FILE MANAGER (DANGEROUS STUFF)


class ReferenceManager:
    def __init__(self, refs_folder: str = REFS_FOLDER):
        self.refs_folder = refs_folder
        self.references = {}
        self.load_all_references()

    def load_all_references(self):
        if not os.path.exists(self.refs_folder):
            os.makedirs(self.refs_folder)
            print(f"Created {self.refs_folder}. Add your reference JSON/TXT files here.")
            return

        existing = set(os.listdir(self.refs_folder))
        desired = set(REFERENCE_FILES)

        for filename in existing - desired:
            filepath = os.path.join(self.refs_folder, filename)
            try:
                if os.path.isfile(filepath):
                    os.remove(filepath)
                    print(f"Removed old reference: {filename}")
            except Exception as e:
                print(f"Could not remove {filename}: {e}")

        for filename in REFERENCE_FILES:
            filepath = os.path.join(self.refs_folder, filename)
            if not os.path.exists(filepath):
                print(f"Warning: reference file missing: {filename}")
                continue

            if filename.endswith('.json'):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        self.references[filename] = json.load(f)
                    print(f"Loaded: {filename}")
                except Exception as e:
                    print(f"Error loading {filename}: {e}")

            elif filename.endswith('.txt'):
                content = None
                for encoding in ("utf-8", "utf-8-sig", "utf-16", "utf-16-le", "utf-16-be", "cp1252", "latin-1"):
                    try:
                        with open(filepath, 'r', encoding=encoding) as f:
                            content = f.read()
                        print(f"Loaded {filename} using encoding {encoding}")
                        break
                    except UnicodeError:
                        continue
                    except Exception as e:
                        print(f"Error loading {filename} with encoding {encoding}: {e}")
                        content = None
                        break

                if content is None:
                    try:
                        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                            content = f.read()
                        print(f"Loaded {filename} with replacement decoding")
                    except Exception as e:
                        print(f"Error loading {filename}: unable to decode file with supported encodings: {e}")
                        continue

                self.references[filename] = self._parse_txt_with_pages(content)
                print(f"Loaded: {filename}")

    def _parse_txt_with_pages(self, content: str) -> Dict:
        lines = content.split('\n')
        result = {"entries": []}

        current_page = 1
        for idx, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            page_match = re.search(r'\[(?:PAGE|p)\s*(\d+)\]', line)
            if page_match:
                current_page = int(page_match.group(1))
                continue

            result["entries"].append({
                "content": line,
                "page": current_page,
                "line": idx + 1
            })

        return result

    def get_reference_context(self, query: str, max_entries: int = 3) -> str:
        """Get context from all reference files (generic)."""
        return self._build_context(self.references.keys(), max_entries)

    def get_specific_reference_context(self, filenames: List[str], max_entries: int = 3) -> str:
        """Get context only from the specified filenames."""
        available = [f for f in filenames if f in self.references]
        if not available:
            return ""
        return self._build_context(available, max_entries)

    def _build_context(self, filenames, max_entries: int) -> str:
        context = "REFERENCE MATERIAL:\n"
        for filename in filenames:
            ref_data = self.references.get(filename)
            if not ref_data:
                continue
            if isinstance(ref_data, dict) and "entries" in ref_data:
                entries = ref_data["entries"][:max_entries]
                context += f"\n[{filename}]\n"
                for entry in entries:
                    context += f"  Page {entry['page']}: {entry['content']}\n"
            elif isinstance(ref_data, dict):
                context += f"\n[{filename}]\n"
                context += json.dumps(ref_data, indent=2)[:500] + "...\n"
        return context if context != "REFERENCE MATERIAL:\n" else ""

    def extract_citations(self, groq_response: str) -> Tuple[str, List[str]]:
        citations = []
        standard = re.findall(r'\[SOURCE:([^\]]+)\]', groq_response)
        citations.extend(standard)
        alt1 = re.findall(r'Source:\s*(Page\s*\d+[^,.]*)', groq_response, re.IGNORECASE)
        citations.extend(alt1)
        alt2 = re.findall(r'\(Page\s*\d+\)', groq_response)
        citations.extend(alt2)
        alt3 = re.findall(r'Page\s*\d+', groq_response)
        
        all_citations = list(set(citations))
        clean_response = groq_response
        for pattern in [r'\[SOURCE:[^\]]+\]', r'Source:\s*Page\s*\d+', r'\(Page\s*\d+\)']:
            clean_response = re.sub(pattern, '', clean_response, flags=re.IGNORECASE)
        clean_response = clean_response.strip()
        
        return clean_response, all_citations


#GRO(Q) API MANAGE

class GroqLessonPlanGenerator:
    def __init__(self, api_key: str, ref_manager: ReferenceManager, diag: bool = False):
        self.client = None
        self.client_type = "none"
        self.api_key = api_key or os.getenv('GROQ_API_KEY') or GROQ_API_KEY
        self.ref_manager = ref_manager
        self.diag = diag
        self.model = "llama-3.3-70b-versatile"
        self.no_method_warning_issued = False

        if Groq is None:
            print(f"{WARNING} Info: Groq SDK not found (groq package). To install: pip install groq")

        if not self.api_key:
            print(f"{WARNING} Info: No GROQ_API_KEY provided. Set the GROQ_API_KEY environment variable or pass a key to the generator.")

        if Groq is not None and self.api_key:
            try:
                self.client = Groq(api_key=self.api_key)
                self.client_type = "groq"
                k = str(self.api_key)
                preview = ("*" * max(0, len(k)-4)) + k[-4:]
                print(f"{CHECK} Groq client initialized (key ending {preview})")
            except Exception as e:
                print(f"{WARNING} Warning: could not initialize Groq client: {e}")
                self.client = None
                self.client_type = "none"

        if self.client_type == "none":
            print(f"{WARNING} Warning: Groq client not initialized or failed. The generator will use a REST fallback or return stubs.")

    def generate_column_content(
        self,
        column_name: str,
        class_info: str,
        subject: str,
        chapter: str,
        period_topic: str,
        prev_content: Dict = None
    ) -> Tuple[str, List[str]]:
        if column_name in ['TOPIC', 'OBJECTIVES', 'LEARNING_OUTCOMES']:
            ref_context = self.ref_manager.get_specific_reference_context([LODOC_FILENAME], max_entries=4)
            cite_instruction = " At the end of your response, include a citation like 'Source: Page X' for each reference used. Use the provided REFERENCE MATERIAL to inform your answer."
        else:
            ref_context = ""
            cite_instruction = " Do not include any references or citations."

        ref_context = self._truncate_context(ref_context, max_chars=500)

        prompt = self._build_prompt(
            column_name,
            class_info,
            subject,
            chapter,
            period_topic,
            ref_context,
            cite_instruction,
            prev_content
        )
        prompt = self._truncate_prompt(prompt, max_chars=1500)

        try:
            print(f"  Generating {column_name}...", end=" ", flush=True)

            if not self.client:
                stub = f"Sample content for {column_name} (Groq client not available)"
                local = DEFAULT_LOCAL_CONTENT.get(column_name, stub)
                print(f"{CHECK} (local fallback used for {column_name})")
                return local, []

            raw_resp = self._request_completion(prompt, column_name)

            if self.diag:
                try:
                    print(f"--- DIAG: raw response type: {type(raw_resp)}")
                    if isinstance(raw_resp, (dict, list)):
                        print(json.dumps(raw_resp, ensure_ascii=False, indent=2)[:4000])
                    else:
                        print(repr(raw_resp)[:2000])
                except Exception as e:
                    print(f"{WARNING} DIAG: could not print raw response: {e}")

            response_text = self._extract_text_from_response(raw_resp)

            if (not response_text) or response_text.strip().startswith('Sample content') or 'API unavailable' in response_text:
                try:
                    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
                    dbg_name = os.path.join(OUTPUT_FOLDER, f'groq_debug_{column_name}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json')
                    with open(dbg_name, 'w', encoding='utf-8') as df:
                        json.dump(raw_resp, df, ensure_ascii=False, indent=2)
                    print(f"\n{WARNING} Debug: wrote raw Groq response to {dbg_name}")
                except Exception:
                    pass

                fallback_text = self._find_longest_string(raw_resp)
                if fallback_text:
                    response_text = fallback_text
                else:
                    local = DEFAULT_LOCAL_CONTENT.get(column_name)
                    if local:
                        response_text = local
                        print(f"{WARNING} Info: using local default for {column_name}")

            content, citations = self.ref_manager.extract_citations(response_text)
            if not citations and column_name in ['TOPIC', 'OBJECTIVES', 'LEARNING_OUTCOMES']:
                page_matches = re.findall(r'Page\s*\d+', response_text, re.IGNORECASE)
                if page_matches:
                    citations = list(set(page_matches))
            print(CHECK)
            return content, citations

        except Exception as e:
            print(f"{CROSS} Error: {e}")
            return f"[Error generating {column_name}]", []

    def _truncate_context(self, context: str, max_chars: int = 500) -> str:
        if len(context) <= max_chars:
            return context
        return context[:max_chars] + "...\n[Context truncated due to length]"

    def _truncate_prompt(self, prompt: str, max_chars: int = 1500) -> str:
        if len(prompt) <= max_chars:
            return prompt
        return prompt[:max_chars] + "\n[Prompt truncated due to length]"

    def _request_completion(self, prompt: str, column_name: str):
        if self.client_type == "groq":
            if hasattr(self.client, 'chat') and hasattr(self.client.chat, 'completions'):
                try:
                    return self.client.chat.completions.create(
                        model=self.model,
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=350
                    )
                except Exception:
                    pass

            if hasattr(self.client, 'request'):
                try:
                    return self.client.request(
                        method='POST',
                        path='/v1/chat/completions',
                        json={
                            'model': self.model,
                            'messages': [{"role": "user", "content": prompt}],
                            'max_tokens': 350
                        }
                    )
                except Exception:
                    pass

            if hasattr(self.client, 'post'):
                try:
                    return self.client.post(
                        '/v1/chat/completions',
                        json={
                            'model': self.model,
                            'messages': [{"role": "user", "content": prompt}],
                            'max_tokens': 350
                        }
                    )
                except Exception:
                    pass

        try:
            resp = self._request_groq_rest(prompt, model=self.model)
            return resp
        except Exception as e:
            if not self.no_method_warning_issued:
                print(f"{WARNING} Warning: Groq SDK methods failed and REST fallback failed: {e}")
                self.no_method_warning_issued = True

        return {'text': f'Sample content for {column_name} (Groq API unavailable).'}

    def _request_groq_rest(self, prompt: str, model: str = None):
        try:
            import urllib.request as _ur
            import urllib.error as _ue
        except Exception:
            raise RuntimeError('urllib not available for Groq REST fallback')

        if not self.api_key:
            raise RuntimeError('No GROQ_API_KEY available for REST call')

        urls = [
            'https://api.groq.ai/v1/chat/completions',
            'https://api.groq.ai/v1/completions'
        ]

        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.api_key}'
        }

        last_err = None
        for url in urls:
            if url.rstrip('/').endswith('/completions') and '/chat/' not in url:
                body_obj = {
                    'model': model or self.model,
                    'prompt': prompt,
                    'max_tokens': 350
                }
            else:
                body_obj = {
                    'model': model or self.model,
                    'messages': [{'role': 'user', 'content': prompt}],
                    'max_tokens': 350
                }

            body = json.dumps(body_obj).encode('utf-8')
            req = _ur.Request(url, data=body, headers=headers, method='POST')
            try:
                with _ur.urlopen(req, timeout=30) as resp:
                    raw = resp.read().decode('utf-8')
                    return json.loads(raw)
            except _ue.HTTPError as e:
                data = e.read().decode('utf-8') if hasattr(e, 'read') else ''
                last_err = f'HTTP {e.code}: {data}'
                continue
            except Exception as e:
                last_err = str(e)
                continue

        raise RuntimeError(f'Groq REST failed: {last_err}')

    def _find_longest_string(self, obj) -> str:
        longest = ''

        def walk(x):
            nonlocal longest
            if x is None:
                return
            if isinstance(x, str):
                if len(x) > len(longest):
                    longest = x
            elif isinstance(x, dict):
                for v in x.values():
                    walk(v)
            elif isinstance(x, list) or isinstance(x, tuple):
                for v in x:
                    walk(v)
            else:
                try:
                    s = str(x)
                    if len(s) > len(longest):
                        longest = s
                except Exception:
                    pass

        try:
            walk(obj)
        except Exception:
            return ''

        return longest

    def _extract_text_from_response(self, resp) -> str:
        if isinstance(resp, str):
            return resp

        if isinstance(resp, dict):
            choices = resp.get('choices')
            if choices and isinstance(choices, list):
                first = choices[0]
                if isinstance(first, dict):
                    if 'text' in first:
                        return first['text']
                    if 'message' in first and isinstance(first['message'], dict):
                        return first['message'].get('content') or ''

            for key in ('text', 'output_text', 'content', 'result'):
                if key in resp and isinstance(resp[key], str):
                    return resp[key]

            results = resp.get('results') or resp.get('output')
            if results and isinstance(results, list) and results:
                first = results[0]
                if isinstance(first, dict):
                    if 'text' in first:
                        return first['text']
                    out = first.get('output') or first.get('content')
                    if out:
                        if isinstance(out, list):
                            for item in out:
                                if isinstance(item, dict) and 'text' in item:
                                    return item['text']

        try:
            if hasattr(resp, 'choices'):
                ch = getattr(resp, 'choices')
                if isinstance(ch, (list, tuple)) and ch:
                    first = ch[0]
                    if isinstance(first, dict) and 'text' in first:
                        return first['text']
                    if hasattr(first, 'text'):
                        return getattr(first, 'text')

            if hasattr(resp, 'content'):
                c = getattr(resp, 'content')
                if isinstance(c, (list, tuple)) and c:
                    first = c[0]
                    if isinstance(first, dict) and 'text' in first:
                        return first['text']
                    if hasattr(first, 'text'):
                        return getattr(first, 'text')

            if hasattr(resp, 'output'):
                out = getattr(resp, 'output')
                if isinstance(out, (list, tuple)) and out:
                    first = out[0]
                    if hasattr(first, 'content'):
                        cont = getattr(first, 'content')
                        if isinstance(cont, (list, tuple)) and cont and hasattr(cont[0], 'text'):
                            return getattr(cont[0], 'text')

            for attr in ('text', 'output_text', 'result'):
                if hasattr(resp, attr):
                    val = getattr(resp, attr)
                    if isinstance(val, str):
                        return val
        except Exception:
            pass

        try:
            if hasattr(resp, 'choices'):
                choices = getattr(resp, 'choices')
                if choices and isinstance(choices, (list, tuple)):
                    first = choices[0]
                    if hasattr(first, 'message') and hasattr(first.message, 'content'):
                        return first.message.content
                    if isinstance(first, dict) and first.get('text'):
                        return first['text']
                    if hasattr(first, 'text'):
                        return getattr(first, 'text')

            if hasattr(resp, 'output'):
                output = getattr(resp, 'output')
                if isinstance(output, (list, tuple)) and output:
                    first = output[0]
                    if hasattr(first, 'content'):
                        cont = getattr(first, 'content')
                        if isinstance(cont, str):
                            return cont
                        if isinstance(cont, dict) and 'text' in cont:
                            return cont['text']
                        if isinstance(cont, (list, tuple)) and cont:
                            item = cont[0]
                            if isinstance(item, str):
                                return item
                            if hasattr(item, 'text'):
                                return getattr(item, 'text')
        except Exception:
            pass

        candidate = ''
        try:
            candidate = str(resp)
        except Exception:
            candidate = ''

        try:
            fl = self._find_longest_string(resp)
            if fl and len(fl) > 20:
                return fl
        except Exception:
            pass

        return candidate or ''

    def _build_prompt(
        self,
        column_name: str,
        class_info: str,
        subject: str,
        chapter: str,
        period_topic: str,
        ref_context: str,
        cite_instruction: str,
        prev_content: Dict
    ) -> str:
        base = f"Class: {class_info}, Subject: {subject}, Chapter: {chapter}, Topic: {period_topic}."
        instructions = {
            "OBJECTIVES": f"{base}\nWrite 2-3 specific, measurable learning objectives using Bloom's taxonomy (Remember, Understand, Apply). Format: 'Student will be able to...' Keep concise (under 80 words).{cite_instruction}\n{ref_context}",
            "LEARNING_OUTCOMES": f"{base}\nList 3-4 measurable learning outcomes using action verbs (identify, classify, explain, describe, calculate). Keep under 100 words.{cite_instruction}\n{ref_context}",
            "TOPIC": f"{base}\nProvide a concise list of subtopics/activities for this period (max 10-12 words).{cite_instruction}\n{ref_context}",
            "MISCONCEPTIONS": f"{base}\nList 3-4 common misconceptions students have. Bullet points. Under 80 words. Do not include references or citations.",
            "PREV_KNOWLEDGE": f"{base}\nList 3-4 key prior knowledge/skills. Under 80 words. Do not include references.",
            "ICE_BREAKER": f"{base}\nSuggest an engaging ice-breaker activity (2-3 min). Under 60 words. No citations.",
            "TEACHING_AIDS": f"{base}\nList teaching aids and materials needed. Under 80 words. No citations.",
            "QUESTION_EXIT_TICKET": f"{base}\nCreate 2-3 exit ticket questions (HOTS level). Under 100 words. No citations.",
            "AI_INTEGRATION": f"{base}\nSuggest 1-2 AI/tech tools. Under 80 words. No citations.",
            "HOMEWORK": f"{base}\nDesign homework. Under 80 words. No citations.",
            "MI_MAPPING_SDG": f"{base}\nIdentify Multiple Intelligences engaged and SDG(s) connected. Under 100 words. No citations.",
            "REFLECTION": f"{base}\nCreate 2-3 reflection questions. Under 80 words. No citations.",
            "CAREER_RS_IDEALS": f"{base}\nConnect to careers and real-world applications. Under 80 words. No citations.",
            "VOLUNTEERING": f"{base}\nSuggest a community volunteering/service activity. Under 80 words. No citations.",
        }
        return instructions.get(column_name, f"Generate content for {column_name} about {period_topic}.")


# YEH H MAJDOOR (DOC CREATOR)

class DocumentPopulator:
    def __init__(self, template_path: str = LP_TEMPLATE):
        self.template_path = template_path
        download_template = DOWNLOAD_TEMPLATE
        if not os.path.exists(download_template):
            download_template = self._find_download_template()

        if download_template and os.path.exists(download_template):
            self.template_path = download_template
            print(f"Using downloaded LP template: {download_template}")
            self.doc = Document(download_template)
        elif os.path.exists(template_path):
            print(f"Using default template: {template_path}")
            self.doc = Document(template_path)
        else:
            self.doc = self._create_default_template()

    def _create_default_template(self) -> Document:
        doc = Document()
        doc.add_heading('LESSON PLAN', 0)
        doc.add_paragraph("Please use the provided LP_formatted.docx template.")
        return doc

    def _find_download_template(self) -> str:
        downloads_dir = os.path.expanduser(r"C:\Users\lenovo\Downloads")
        if not os.path.isdir(downloads_dir):
            return ""

        candidates = []
        for filename in os.listdir(downloads_dir):
            lower = filename.lower()
            if filename.lower().endswith('.docx') and ('lp' in lower or 'lesson plan' in lower or 'formatted' in lower):
                candidates.append(os.path.join(downloads_dir, filename))

        return candidates[0] if candidates else ""

    def populate_lp(
        self,
        user_input: Dict,
        generated_content: Dict,
        citations: Dict
    ) -> str:
        replacements = self._build_replacements(user_input, generated_content)
        self._replace_placeholders_in_doc(replacements)
        output_path = self._get_output_path(user_input)
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            self.doc.save(output_path)
        except Exception as e:
            print(f"{CROSS} Failed saving DOCX: {e}")
            raise
        return output_path

    def _build_replacements(self, user_input: Dict, generated_content: Dict) -> Dict[str, str]:
        now_date = datetime.now().strftime('%d-%m-%Y')
        start_date_str = user_input.get('start_date', now_date)
        try:
            start_date = datetime.strptime(start_date_str, '%d-%m-%Y')
        except ValueError:
            start_date = datetime.now()
            start_date_str = start_date.strftime('%d-%m-%Y')

        total_periods = int(user_input.get('total_periods', 3))
        duration = user_input.get('duration', '40 min')

        period_dates = []
        for i in range(total_periods):
            date_i = start_date + timedelta(days=i)
            period_dates.append({
                'date': date_i.strftime('%d-%m-%Y'),
                'day': date_i.strftime('%A')
            })

        replacements = {
            '{{TEACHER}}': user_input.get('teacher', ''),
            '{{CLASS_SECTION}}': user_input.get('class', ''),
            '{{SUBJECT}}': user_input.get('subject', ''),
            '{{CHAPTER}}': user_input.get('chapter', ''),
            '{{CHAPTER_TOPIC}}': user_input.get('chapter', ''),
            '{{DATE}}': f"{start_date_str} to {period_dates[-1]['date']}" if total_periods > 1 else start_date_str,
            '{{TOTAL_PERIODS}}': str(total_periods),
            '{{DURATION}}': duration,
        }

        for p in range(1, total_periods + 1):
            date_info = period_dates[p-1]
            replacements[f'{{{{DATE_{p}}}}}'] = date_info['date']
            replacements[f'{{{{DAY_{p}}}}}'] = date_info['day']
            replacements[f'{{{{PERIOD_{p}}}}}'] = f"Period {p}"
            replacements[f'{{{{TOPIC_{p}}}}}'] = generated_content.get(f'TOPIC_{p}', user_input.get(f'topic_p{p}', ''))
            replacements[f'{{{{LO_{p}}}}}'] = generated_content.get(f'LO_{p}', '')
            replacements[f'{{{{OUTCOME_{p}}}}}'] = generated_content.get(f'OUTCOME_{p}', '')
            replacements[f'{{{{MISCONC_{p}}}}}'] = generated_content.get(f'MISCONC_{p}', '')
            replacements[f'{{{{PRIOR_{p}}}}}'] = generated_content.get(f'PRIOR_{p}', '')
            replacements[f'{{{{ICEBREAK_{p}}}}}'] = generated_content.get(f'ICEBREAK_{p}', '')
            replacements[f'{{{{AIDS_{p}}}}}'] = generated_content.get(f'AIDS_{p}', '')
            replacements[f'{{{{EXIT_{p}}}}}'] = generated_content.get(f'EXIT_{p}', '')
            replacements[f'{{{{AI_{p}}}}}'] = generated_content.get(f'AI_{p}', '')
            replacements[f'{{{{HW_{p}}}}}'] = generated_content.get(f'HW_{p}', '')
            replacements[f'{{{{MI_{p}}}}}'] = generated_content.get(f'MI_{p}', '')
            replacements[f'{{{{SDG_{p}}}}}'] = generated_content.get(f'SDG_{p}', '')
            replacements[f'{{{{VOLUNTEER_{p}}}}}'] = generated_content.get(f'VOLUNTEER_{p}', '')
            replacements[f'{{{{REMARKS_{p}}}}}'] = generated_content.get(f'REMARKS_{p}', '')

        return replacements

    def _replace_placeholders_in_doc(self, replacements: Dict[str, str]):
        for paragraph in self.doc.paragraphs:
            self._replace_placeholders_in_paragraph(paragraph, replacements)
        for table in self.doc.tables:
            self._replace_placeholders_in_table(table, replacements)

    def _replace_placeholders_in_paragraph(self, paragraph, replacements: Dict[str, str]):
        text = ''.join(run.text for run in paragraph.runs)
        new_text = text
        for placeholder, value in replacements.items():
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, value)
        if new_text != text:
            for run in paragraph.runs:
                run.text = ''
            paragraph.add_run(new_text)

    def _replace_placeholders_in_table(self, table, replacements: Dict[str, str]):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, replacements)

    def _get_output_path(self, user_input: Dict) -> str:
        if not os.path.exists(OUTPUT_FOLDER):
            os.makedirs(OUTPUT_FOLDER)
        class_name = user_input.get('class', 'X').replace(' ', '_')
        chapter_name = user_input.get('chapter', 'Lesson').replace(' ', '_')[:20]
        teacher_name = user_input.get('teacher', 'Teacher').replace(' ', '_')[:10]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"LP_{class_name}_{chapter_name}_{teacher_name}_{timestamp}.docx"
        return os.path.join(OUTPUT_FOLDER, filename)


#BIHARI ORCHESTRATOR

class LessonPlanOrchestrator:
    def __init__(self, diag: bool = False, cli_args: Dict = None):
        self.diag = diag
        self.cli_args = cli_args or {}
        self.ref_manager = ReferenceManager(REFS_FOLDER)
        self.groq_gen = GroqLessonPlanGenerator(GROQ_API_KEY, self.ref_manager, diag=diag)
        self.doc_pop = DocumentPopulator()

    def get_user_input(self) -> Dict:
        # If CLI args provided, use them
        if self.cli_args:
            print("\n" + "="*60)
            print("LESSON PLAN GENERATOR - Using command-line arguments")
            print("="*60)
            
            # Ensure all expected fields exist
            user_data = {
                'teacher': self.cli_args.get('teacher', 'Teacher'),
                'class': self.cli_args.get('class', 'X A'),
                'subject': self.cli_args.get('subject', 'Biology'),
                'chapter': self.cli_args.get('chapter', 'Nutrition'),
                'start_date': self.cli_args.get('start_date', datetime.now().strftime('%d-%m-%Y')),
                'duration': self.cli_args.get('duration', '40 min'),
                'total_periods': self.cli_args.get('total_periods', '3'),
            }
            
            # Get topics from CLI args
            total = int(user_data['total_periods'])
            for p in range(1, total + 1):
                user_data[f'topic_p{p}'] = self.cli_args.get(f'topic_p{p}', f"Topic {p}")
            
            print(f"Teacher: {user_data['teacher']}")
            print(f"Class: {user_data['class']}")
            print(f"Subject: {user_data['subject']}")
            print(f"Chapter: {user_data['chapter']}")
            print(f"Start Date: {user_data['start_date']}")
            print(f"Total Periods: {user_data['total_periods']}")
            for p in range(1, total + 1):
                print(f"  Period {p}: {user_data[f'topic_p{p}']}")
            
            return user_data
        
        # Interactive mode
        print("\n" + "="*60)
        print("LESSON PLAN GENERATOR - Input")
        print("="*60)

        user_data = {
            'teacher': input("Teacher Name: ").strip() or "Teacher",
            'class': input("Class & Section (e.g., X A): ").strip() or "X A",
            'subject': input("Subject (e.g., Biology): ").strip() or "Biology",
            'chapter': input("Chapter (e.g., Nutrition): ").strip() or "Nutrition",
            'start_date': input("Start Date (DD-MM-YYYY): ").strip() or datetime.now().strftime('%d-%m-%Y'),
            'duration': input("Duration per period (e.g., 40 min): ").strip() or "40 min",
            'total_periods': input("Total number of periods (default 3): ").strip() or "3",
        }

        total = int(user_data['total_periods'])
        for p in range(1, total + 1):
            topic = input(f"Topic for Period {p}: ").strip() or f"Topic {p}"
            user_data[f'topic_p{p}'] = topic

        return user_data

    def generate_lesson_plan(self, user_input: Dict):
        print("\n" + "="*60)
        print(f"Generating LP for {user_input['class']} - {user_input['chapter']}")
        print("="*60 + "\n")

        generated_content = {}
        citations_data = {}
        total_periods = int(user_input.get('total_periods', 3))

        for p in range(1, total_periods + 1):
            period_topic = user_input.get(f'topic_p{p}', f"Topic {p}")
            generated_content[f'TOPIC_{p}'] = period_topic

            mapping = {
                'LO': 'OBJECTIVES',
                'OUTCOME': 'LEARNING_OUTCOMES',
                'MISCONC': 'MISCONCEPTIONS',
                'PRIOR': 'PREV_KNOWLEDGE',
                'ICEBREAK': 'ICE_BREAKER',
                'AIDS': 'TEACHING_AIDS',
                'EXIT': 'QUESTION_EXIT_TICKET',
                'AI': 'AI_INTEGRATION',
                'HW': 'HOMEWORK',
                'MI': 'MI_MAPPING_SDG',
                'SDG': 'MI_MAPPING_SDG',
                'VOLUNTEER': 'VOLUNTEERING',
                'REMARKS': 'REFLECTION'
            }

            for key, col_name in mapping.items():
                if key == 'SDG' and f'MI_{p}' in generated_content:
                    generated_content[f'SDG_{p}'] = generated_content[f'MI_{p}']
                    continue

                content, citations = self.groq_gen.generate_column_content(
                    col_name,
                    user_input['class'],
                    user_input['subject'],
                    user_input['chapter'],
                    period_topic,
                    prev_content=generated_content
                )
                generated_content[f'{key}_{p}'] = content
                if citations:
                    citations_data[f'{key}_{p}'] = citations

            if f'MI_{p}' in generated_content and f'SDG_{p}' not in generated_content:
                generated_content[f'SDG_{p}'] = generated_content[f'MI_{p}']

        print("\nPopulating Word document...")
        output_path = self.doc_pop.populate_lp(user_input, generated_content, citations_data)
        print(f"\n{CHECK} LP saved: {output_path}")
        return output_path

    def run(self):
        try:
            user_input = self.get_user_input()
            self.generate_lesson_plan(user_input)
            print("\nDone! Your lesson plan is ready.")
        except KeyboardInterrupt:
            print("\n\nExecution cancelled.")
        except Exception as e:
            print(f"\n{CROSS} Error: {e}")
            import traceback
            traceback.print_exc()



# ENTRY POINT H YEH


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Lesson Plan Generator')
    parser.add_argument('--diag', action='store_true', help='Print diagnostic Groq responses to console')
    
    # Command-line arguments for all inputs
    parser.add_argument('--teacher', type=str, help='Teacher name')
    parser.add_argument('--class', dest='class_name', type=str, help='Class and section')
    parser.add_argument('--subject', type=str, help='Subject name')
    parser.add_argument('--chapter', type=str, help='Chapter name')
    parser.add_argument('--start_date', type=str, help='Start date (DD-MM-YYYY)')
    parser.add_argument('--duration', type=str, help='Duration per period')
    parser.add_argument('--total_periods', type=str, help='Total number of periods')
    
    # Dynamic topic arguments (we'll parse them manually to support any number)
    args, unknown = parser.parse_known_args()
    
    # Parse topic arguments (--topic_p1, --topic_p2, etc.)
    topic_args = {}
    for arg in unknown:
        if arg.startswith('--topic_p'):
            # The next argument is the value
            idx = unknown.index(arg)
            if idx + 1 < len(unknown) and not unknown[idx + 1].startswith('--'):
                topic_key = arg[2:]  # Remove '--'
                topic_args[topic_key] = unknown[idx + 1]
    
    # Build cli_args dict for the orchestrator
    cli_args = {
        'teacher': getattr(args, 'teacher', None),
        'class': getattr(args, 'class_name', None),
        'subject': args.subject,
        'chapter': args.chapter,
        'start_date': args.start_date,
        'duration': args.duration,
        'total_periods': args.total_periods,
        **topic_args
    }
    
    # Only pass cli_args if any were provided
    if any(v is not None for v in cli_args.values()):
        orchestrator = LessonPlanOrchestrator(diag=args.diag, cli_args=cli_args)
    else:
        orchestrator = LessonPlanOrchestrator(diag=args.diag)
    
    orchestrator.run()
